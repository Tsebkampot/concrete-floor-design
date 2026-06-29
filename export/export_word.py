"""Word 计算书导出工具。"""

from __future__ import annotations

from io import BytesIO

import pandas as pd
from docx import Document


def _add_dataframe_table(document: Document, df: pd.DataFrame) -> None:
    """将 DataFrame 写入 Word 表格。"""
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for cell, column in zip(table.rows[0].cells, df.columns):
        cell.text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)


def build_word_report(title: str, tables: dict[str, pd.DataFrame], notes: list[str] | None = None) -> bytes:
    """用 python-docx 和 BytesIO 生成 docx 半自动计算书。"""
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph("课程名称：《水工钢筋混凝土》")
    document.add_paragraph("题目：整体式单向板肋形楼盖设计与辅助计算程序开发")
    document.add_heading("人工复核说明", level=1)
    document.add_paragraph("本计算书由程序半自动生成，结果不能替代课程设计手算、教材系数表和教师要求的人工复核。")

    if notes:
        document.add_heading("已知不足和适用范围", level=1)
        for note in notes:
            document.add_paragraph(str(note), style="List Bullet")

    for heading, df in tables.items():
        document.add_heading(str(heading), level=1)
        _add_dataframe_table(document, df)

    output = BytesIO()
    document.save(output)
    return output.getvalue()
